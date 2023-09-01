from docx import Document
from docx.shared import Pt
import re
import os

# Получаем текущую рабочую директорию
working_dir = os.getcwd()

# Получаем путь к родительской директории
parent_dir = os.path.dirname(working_dir)
article_path = os.path.join(parent_dir, "doc", "article.txt")
docx_path = os.path.join(parent_dir, "doc", "article.docx")

# Открываем файл .txt и читаем его содержимое
with open(article_path, 'r', encoding='utf-8') as txt_file:
    txt_content = txt_file.read()

# Убираем лишние пробелы между строками (3 и более заменяем на 2)
txt_content = re.sub(r'(\n\s{3,})', '\n\n  ', txt_content)

# Создаем новый документ .docx и добавляем текст
doc = Document()
first_paragraph = doc.add_paragraph(txt_content.split('\n', 1)[0])
run = first_paragraph.runs[0]
run.font.size = Pt(18)  # 1.5rem 
run.font.name = 'Fira Sans'

# Добавляем оставшийся текст как обычно
doc.add_paragraph(txt_content[len(first_paragraph.text) + 1:])


# Сохраняем документ в файл .docx
doc.save(docx_path)
